import os
import time
import re
from google import genai 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# --- Basic Configurations ---
API_KEY = "your-API_KEY" 
client = genai.Client(api_key=API_KEY) 

def is_arabic(text):
    """Checks if the line starts with Arabic characters"""
    if not text: return False
    # Check the first alphabetic character in the string
    first_alpha = re.search(r'[a-zA-Z\u0600-\u06FF]', text)
    if first_alpha:
        char = first_alpha.group()
        return '\u0600' <= char <= '\u06FF'
    return True # Default to True (Arabic) if no alphabetic characters found

def compress_audio(input_path, output_path):
    """Compresses audio using ffmpeg to reduce size"""
    if not os.path.exists(input_path):
        return False, "File does not exist"
    cmd = f'ffmpeg -y -i "{input_path}" -ab 64k -ac 1 "{output_path}"'
    return (os.system(cmd) == 0), output_path

def create_formatted_doc(text_content, output_file):
    """Automatically formats the document based on text language without a template"""
    if not text_content:
        return False, "No content found"

    doc = Document() # Create a brand new document
    
    # Formatting Colors
    HEADING_COLOR = RGBColor(0, 112, 192) # Blue color for headers
    
    # Clean text from markdown-style stars (*)
    text_content = text_content.replace('*', '')

    lines = text_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = doc.add_paragraph()
        
        # Detect language for alignment and direction
        if is_arabic(line):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.rtl = True
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.rtl = False

        # If it's a heading (starts with #)
        if line.startswith('#'):
            clean_line = line.replace('#', '').strip()
            run = p.add_run(clean_line)
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = HEADING_COLOR
        else:
            run = p.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
        # Set RTL for the specific Run if the line is Arabic
        if is_arabic(line):
            run.font.rtl = True

    doc.save(output_file)
    return True, output_file

def process_audio_to_docx(audio_path):
    """Main function to upload audio, transcribe using AI, and save to Word"""
    uploaded_file = None
    try:
        print(f"[*] Uploading and analyzing file...")
        uploaded_file = client.files.upload(
            file=audio_path,
            config={'mime_type': 'audio/mpeg'}
        )
        
        # Wait for processing to finish
        while uploaded_file.state.name == "PROCESSING":
            time.sleep(2)
            uploaded_file = client.files.get(name=uploaded_file.name)

        system_prompt = (
            "You are an expert transcriptionist for all languages. "
            "Accurately and comprehensively transcribe the audio content. "
            "Organize the text using headings (#) for new paragraphs or sections. "
            "If the speaker talks in Arabic, write in Arabic. If they speak English, write in English. "
            "Do not add any side comments; start directly with the content."
        )

        response = client.models.generate_content(
            model="gemini-2.5-flash", # Updated to current latest model
            contents=[system_prompt, uploaded_file]
        )
        
        full_text = response.text if response.text else ""
        
        if full_text:
            output_name = f"transcription_{int(time.time())}.docx"
            success, result = create_formatted_doc(full_text, output_name)
            return success, result
        
        return False, "The AI response was empty"

    except Exception as e:
        return False, str(e)
    finally:
        # Cleanup: Delete the file from AI storage
        if uploaded_file:
            client.files.delete(name=uploaded_file.name)